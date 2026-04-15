"""
innova-converter — Google Cloud Run
Recibe datos JSON del despacho, llena el Word original,
convierte a PDF con LibreOffice y devuelve el PDF en base64.
"""
import os
import base64
import subprocess
import tempfile
import json
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
CORS(app, origins=["https://alexgf004-maker.github.io"])

TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Entrega_de_materiales_OTC_2026.docx")

SAP_ROW = {
    "221477":2,"213719":3,"328541":4,"352453":5,"352460":6,
    "352461":7,"352462":8,"353112":9,"354045":10,"354549":11,
    "200129":12,"355518":13,"338362":14,"219359":15,
    "328560":17,"243940":18,"337775":19,"337776":20,"337777":21,
    "210525":22,"212720":23,"214221":24,"301560":25,"245979":26,
    "355070":27,"244569":28,"353992":29,"211373":30,"211375":31,
    "353121":32,"355064":33,"338357":34,"353099":35,"353110":36,
    "353088":37,"200468":39,"200472":40,"200469":41,"200473":42,
    "213410":43,"214726":44,"214727":45,"352463":46,
    "219527":48,"221062":49,"200367":50,"282485":51,"350560":52,
    "212896":53,"350564":54,"221472":56,"200413":57,"211829":58,
    "213340":59,"222315":60,"353730":61,"338363":62,"338361":63,"338360":64,
}

def fill_header_cell(cell, value):
    if not value:
        return
    para = cell.paragraphs[0]
    if para.runs:
        last_run = para.runs[-1]
        if '\t' in last_run.text:
            last_run.text = last_run.text.replace('\t', str(value))
            last_run.bold = True
            return
    run = para.add_run(str(value))
    run.bold = True

def fill_quantity_cell(cell, value):
    if not value:
        return
    para = cell.paragraphs[0]
    font_size = next((r.font.size for r in para.runs if r.font.size), None)
    run = para.add_run(str(value))
    run.bold = True
    run.font.size = font_size if font_size else Pt(7.5)

def fill_template(data):
    doc = Document(TEMPLATE)
    t0   = doc.tables[0]
    rows = t0.rows

    campos = {
        (3, 1): data.get("usuarioResponsable",    ""),
        (4, 0): data.get("empresaContratista",    ""),
        (4, 1): data.get("instaladorResponsable", ""),
        (5, 0): data.get("entregadoPor",          ""),
        (6, 1): data.get("placaVehiculo",         ""),
        (7, 0): data.get("fechaSolicitud",        ""),
        (7, 1): data.get("fechaEntrega",          ""),
    }
    for (fila, col), valor in campos.items():
        if valor:
            try:
                fill_header_cell(rows[fila].cells[col], valor)
            except IndexError:
                pass

    t1_rows = doc.tables[1].rows
    for item in data.get("items", []):
        sap     = str(item.get("sapCode", "")).strip()
        cant    = item.get("cantidad", "")
        row_idx = SAP_ROW.get(sap)
        if row_idx is not None and cant:
            try:
                fill_quantity_cell(t1_rows[row_idx].cells[3], str(cant))
            except IndexError:
                pass

    return doc

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

@app.route('/generar-pdf', methods=['POST'])
def generar_pdf():
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Falta el cuerpo JSON'}), 400

        if not os.path.exists(TEMPLATE):
            return jsonify({'error': 'Plantilla Word no encontrada en el servidor'}), 500

        with tempfile.TemporaryDirectory() as tmpdir:
            # 1. Llenar el Word
            doc      = fill_template(data)
            docx_path = os.path.join(tmpdir, 'despacho.docx')
            pdf_path  = os.path.join(tmpdir, 'despacho.pdf')
            doc.save(docx_path)

            # 2. Convertir a PDF con LibreOffice
            result = subprocess.run([
                'libreoffice', '--headless', '--norestore',
                '--convert-to', 'pdf',
                '--outdir', tmpdir, docx_path
            ], capture_output=True, text=True, timeout=45,
               env={**os.environ, 'HOME': tmpdir})

            if result.returncode != 0:
                return jsonify({'error': f'Conversión fallida: {result.stderr}'}), 500

            if not os.path.exists(pdf_path):
                return jsonify({'error': 'PDF no generado'}), 500

            # 3. Devolver PDF en base64
            with open(pdf_path, 'rb') as f:
                pdf_b64 = base64.b64encode(f.read()).decode('utf-8')

        return jsonify({'pdf_base64': pdf_b64})

    except subprocess.TimeoutExpired:
        return jsonify({'error': 'Timeout — conversión tardó más de 45s'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
